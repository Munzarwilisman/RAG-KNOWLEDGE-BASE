"""
╔══════════════════════════════════════════════════════════════╗
║         DIGIT-OPS RAG — Knowledge Base PLTU                 ║
║         Tim 2-10 orang | Supabase PostgreSQL                ║
║         Bilingual ID+EN | Claude Vision OCR                 ║
╚══════════════════════════════════════════════════════════════╝
"""

import streamlit as st

# ── Page config HARUS paling atas ─────────────────────────────
st.set_page_config(
    page_title="DIGIT-OPS Knowledge Base",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

import os, json, hashlib, re, time, datetime, pickle, base64
import numpy as np
import pandas as pd
from pathlib import Path
from io import BytesIO
from typing import Optional

# ══════════════════════════════════════════════════════════════
# KONSTANTA
# ══════════════════════════════════════════════════════════════
APP_VERSION    = "2.0.0"
INDEX_DIR      = Path("rag_index")
INDEX_DIR.mkdir(exist_ok=True)

CHUNK_SIZE     = 1200
CHUNK_OVERLAP  = 200
TOP_K          = 8
EMBED_MODEL    = "paraphrase-multilingual-MiniLM-L12-v2"
IMG_THRESHOLD  = 120   # karakter min sebelum OCR dijalankan
OCR_DPI        = 150

DOC_CATEGORIES = [
    "Manual Book — Boiler CFB",
    "Manual Book — Steam Turbine",
    "Manual Book — Generator",
    "Manual Book — BOP & Auxiliaries",
    "Instrumen & Kontrol",
    "SOP & Prosedur Operasi",
    "Maintenance & Overhaul",
    "Jurnal & Referensi Teknis",
    "Standar (EPRI / ASME / IEC)",
    "Keselamatan (K3 & HSE)",
    "Data Sheet & Spesifikasi",
    "Laporan & Analisis",
    "Lain-lain",
]

# ══════════════════════════════════════════════════════════════
# DATABASE — SUPABASE POSTGRESQL
# ══════════════════════════════════════════════════════════════

@st.cache_resource(show_spinner=False)
def get_db():
    """Koneksi PostgreSQL via DATABASE_URL. Di-cache agar tidak reconnect tiap rerun."""
    try:
        import psycopg2
        import psycopg2.extras
        url  = st.secrets["DATABASE_URL"]
        conn = psycopg2.connect(url, sslmode="require")
        conn.autocommit = False
        _init_tables(conn)
        return conn
    except Exception as e:
        return None


def _init_tables(conn):
    """Buat tabel jika belum ada."""
    sql = """
    CREATE TABLE IF NOT EXISTS rag_files (
        filename    TEXT PRIMARY KEY,
        data        BYTEA NOT NULL,
        updated_at  TIMESTAMPTZ DEFAULT NOW()
    );
    CREATE TABLE IF NOT EXISTS rag_documents (
        doc_id      TEXT PRIMARY KEY,
        name        TEXT NOT NULL,
        category    TEXT,
        description TEXT,
        pages       INT,
        n_chunks    INT,
        n_text      INT,
        n_ocr       INT,
        n_error     INT,
        file_hash   TEXT UNIQUE,
        ocr_engine  TEXT,
        uploaded_by TEXT,
        created_at  TIMESTAMPTZ DEFAULT NOW(),
        updated_at  TIMESTAMPTZ DEFAULT NOW()
    );
    CREATE TABLE IF NOT EXISTS rag_chunks (
        id          SERIAL PRIMARY KEY,
        doc_id      TEXT NOT NULL REFERENCES rag_documents(doc_id) ON DELETE CASCADE,
        page_num    INT,
        page_type   TEXT,
        category    TEXT,
        chunk_text  TEXT NOT NULL,
        created_at  TIMESTAMPTZ DEFAULT NOW()
    );
    CREATE INDEX IF NOT EXISTS idx_chunks_doc ON rag_chunks(doc_id);
    CREATE INDEX IF NOT EXISTS idx_chunks_cat ON rag_chunks(category);
    """
    try:
        cur = conn.cursor()
        cur.execute(sql)
        conn.commit()
        cur.close()
    except Exception as e:
        conn.rollback()


def db_push_index(index_dir: Path) -> tuple[bool, str]:
    """Upload FAISS index files ke PostgreSQL."""
    conn = get_db()
    if not conn:
        return False, "Database tidak terhubung"
    try:
        import psycopg2
        cur = conn.cursor()
        pushed = []
        for fname in ["chunks.pkl", "faiss.index"]:
            fpath = index_dir / fname
            if fpath.exists():
                with open(fpath, "rb") as f:
                    data = f.read()
                cur.execute("""
                    INSERT INTO rag_files (filename, data, updated_at)
                    VALUES (%s, %s, NOW())
                    ON CONFLICT (filename) DO UPDATE
                    SET data = EXCLUDED.data, updated_at = NOW()
                """, (fname, psycopg2.Binary(data)))
                pushed.append(fname)
        conn.commit()
        cur.close()
        return True, f"✅ Pushed: {', '.join(pushed)}"
    except Exception as e:
        conn.rollback()
        return False, f"❌ Push gagal: {e}"


def db_pull_index(index_dir: Path) -> tuple[bool, str]:
    """Download FAISS index files dari PostgreSQL."""
    conn = get_db()
    if not conn:
        return False, "Database tidak terhubung"
    try:
        cur = conn.cursor()
        pulled = []
        for fname in ["chunks.pkl", "faiss.index"]:
            cur.execute("SELECT data FROM rag_files WHERE filename = %s", (fname,))
            row = cur.fetchone()
            if row:
                index_dir.mkdir(exist_ok=True)
                with open(index_dir / fname, "wb") as f:
                    f.write(bytes(row[0]))
                pulled.append(fname)
        cur.close()
        if not pulled:
            return False, "Index belum ada (upload dokumen pertama kali)"
        return True, f"✅ Pulled: {', '.join(pulled)}"
    except Exception as e:
        return False, f"❌ Pull gagal: {e}"


def db_get_all_docs() -> list[dict]:
    """Ambil semua metadata dokumen dari PostgreSQL."""
    conn = get_db()
    if not conn:
        return []
    try:
        import psycopg2.extras
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT * FROM rag_documents ORDER BY created_at DESC")
        rows = [dict(r) for r in cur.fetchall()]
        cur.close()
        return rows
    except Exception:
        return []


def db_save_doc_meta(doc_id: str, meta: dict) -> bool:
    """Simpan metadata dokumen ke PostgreSQL."""
    conn = get_db()
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO rag_documents
              (doc_id, name, category, description, pages, n_chunks,
               n_text, n_ocr, n_error, file_hash, ocr_engine,
               uploaded_by, created_at, updated_at)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,NOW(),NOW())
            ON CONFLICT (doc_id) DO UPDATE SET
              name=EXCLUDED.name, category=EXCLUDED.category,
              description=EXCLUDED.description, pages=EXCLUDED.pages,
              n_chunks=EXCLUDED.n_chunks, n_text=EXCLUDED.n_text,
              n_ocr=EXCLUDED.n_ocr, n_error=EXCLUDED.n_error,
              ocr_engine=EXCLUDED.ocr_engine,
              uploaded_by=EXCLUDED.uploaded_by,
              updated_at=NOW()
        """, (
            doc_id, meta["name"], meta["category"], meta.get("description",""),
            meta.get("pages",0), meta.get("n_chunks",0),
            meta.get("n_text",0), meta.get("n_ocr",0), meta.get("n_error",0),
            meta.get("hash",""), meta.get("ocr_engine","claude"),
            meta.get("uploaded_by",""),
        ))
        conn.commit()
        cur.close()
        return True
    except Exception as e:
        conn.rollback()
        return False


def db_delete_doc(doc_id: str) -> bool:
    """Hapus dokumen dari PostgreSQL (CASCADE hapus chunks juga)."""
    conn = get_db()
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM rag_documents WHERE doc_id = %s", (doc_id,))
        conn.commit()
        cur.close()
        return True
    except Exception as e:
        conn.rollback()
        return False


def db_hash_exists(file_hash: str) -> bool:
    """Cek apakah file sudah pernah diupload (by hash)."""
    conn = get_db()
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute("SELECT 1 FROM rag_documents WHERE file_hash = %s", (file_hash,))
        exists = cur.fetchone() is not None
        cur.close()
        return exists
    except Exception:
        return False


def is_db_connected() -> bool:
    """Cek apakah koneksi database tersedia."""
    try:
        conn = get_db()
        if not conn:
            return False
        cur = conn.cursor()
        cur.execute("SELECT 1")
        cur.close()
        return True
    except Exception:
        return False


# ══════════════════════════════════════════════════════════════
# OCR — CLAUDE VISION
# ══════════════════════════════════════════════════════════════

def ocr_claude(img_bytes: bytes, page_num: int, doc_name: str) -> str:
    """Claude Vision OCR — akurat untuk tabel teknis bilingual."""
    try:
        import anthropic
        client  = anthropic.Anthropic(api_key=st.secrets["anthropic"]["api_key"])
        img_b64 = base64.b64encode(img_bytes).decode()
        prompt  = f"""Kamu adalah OCR engine untuk dokumen teknis PLTU/Power Plant.
Dokumen: {doc_name} | Halaman: {page_num}

INSTRUKSI:
1. Ekstrak SEMUA teks dari halaman ini secara lengkap dan akurat.
2. Untuk TABEL PARAMETER — format setiap baris:
   [Nama Parameter] | [Unit] | [Nilai Normal] | [Alarm High] | [Alarm Low] | [Trip/Interlock]
   WAJIB: nama parameter HARUS ada di setiap baris, jangan pisahkan nama dari nilainya.
3. Untuk tabel clearance/alignment:
   [Komponen] | [Symbol] | [Nilai mm] | [Toleransi]
4. Teks bilingual Mandarin+Inggris → ekstrak KEDUANYA.
5. Angka teknis harus AKURAT: °C, MPa, kPa, bar, mm, rpm, MW.
6. Jika halaman kosong atau tidak ada teks, tulis: [Halaman kosong]

Ekstrak seluruh isi halaman:"""
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=3000,
            messages=[{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64",
                 "media_type": "image/jpeg", "data": img_b64}},
                {"type": "text", "text": prompt},
            ]}],
        )
        return resp.content[0].text
    except Exception as e:
        return f"[OCR Error hal {page_num}: {e}]"


# ══════════════════════════════════════════════════════════════
# PDF PROCESSING
# ══════════════════════════════════════════════════════════════

def pdf_page_count(pdf_bytes: bytes) -> int:
    for lib in ["fitz", "pypdf", "PyPDF2"]:
        try:
            if lib == "fitz":
                import fitz
                return len(fitz.open(stream=pdf_bytes, filetype="pdf"))
            elif lib == "pypdf":
                import pypdf
                return len(pypdf.PdfReader(BytesIO(pdf_bytes)).pages)
            elif lib == "PyPDF2":
                import PyPDF2
                return len(PyPDF2.PdfReader(BytesIO(pdf_bytes)).pages)
        except Exception:
            continue
    return 0


def pdf_extract_text(pdf_bytes: bytes, page_num: int) -> str:
    for lib in ["fitz", "pypdf", "PyPDF2"]:
        try:
            if lib == "fitz":
                import fitz
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                return doc[page_num - 1].get_text().strip()
            elif lib == "pypdf":
                import pypdf
                return (pypdf.PdfReader(BytesIO(pdf_bytes))
                        .pages[page_num - 1].extract_text() or "").strip()
            elif lib == "PyPDF2":
                import PyPDF2
                return (PyPDF2.PdfReader(BytesIO(pdf_bytes))
                        .pages[page_num - 1].extract_text() or "").strip()
        except Exception:
            continue
    return ""


def pdf_rasterize(pdf_bytes: bytes, page_num: int, dpi: int = OCR_DPI) -> Optional[bytes]:
    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = doc[page_num - 1].get_pixmap(matrix=mat, alpha=False)
        return pix.tobytes("jpeg")
    except Exception:
        pass
    try:
        import pypdfium2 as pdfium
        doc  = pdfium.PdfDocument(pdf_bytes)
        page = doc[page_num - 1]
        bm   = page.render(scale=dpi / 72)
        buf  = BytesIO()
        bm.to_pil().save(buf, format="JPEG", quality=85)
        return buf.getvalue()
    except Exception:
        pass
    return None


def process_pdf(pdf_bytes: bytes, filename: str, category: str,
                progress_cb=None) -> list[dict]:
    """Proses PDF dengan hybrid text + OCR."""
    pages    = []
    n_pages  = pdf_page_count(pdf_bytes)
    if n_pages == 0:
        return pages

    n_text = n_ocr = n_err = 0
    for pg in range(1, n_pages + 1):
        if progress_cb:
            progress_cb(pg, n_pages, n_text, n_ocr)

        text = pdf_extract_text(pdf_bytes, pg)

        if len(text) >= IMG_THRESHOLD:
            pages.append({"text": text, "page": pg,
                          "type": "text", "source": filename, "category": category})
            n_text += 1
        else:
            img = pdf_rasterize(pdf_bytes, pg)
            if img:
                ocr_text = ocr_claude(img, pg, filename)
                if "[OCR Error" not in ocr_text:
                    combined = (text + "\n\n" + ocr_text).strip() if text else ocr_text
                    pages.append({"text": combined, "page": pg,
                                  "type": "claude_ocr", "source": filename, "category": category})
                    n_ocr += 1
                else:
                    if text:
                        pages.append({"text": text, "page": pg,
                                      "type": "text", "source": filename, "category": category})
                        n_text += 1
                    else:
                        n_err += 1
            else:
                if text:
                    pages.append({"text": text, "page": pg,
                                  "type": "text", "source": filename, "category": category})
                    n_text += 1
                else:
                    n_err += 1
    return pages


def process_docx(file_bytes: bytes, filename: str, category: str) -> list[dict]:
    pages = []
    try:
        from docx import Document
        doc   = Document(BytesIO(file_bytes))
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Ambil tabel juga
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        # Chunk per 500 words
        words = " ".join(paras).split()
        for i in range(0, len(words), 500):
            chunk = " ".join(words[i:i + 500])
            if chunk:
                pages.append({"text": chunk, "page": i // 500 + 1,
                               "type": "text", "source": filename, "category": category})
    except Exception as e:
        pages.append({"text": f"Error: {e}", "page": 1,
                      "type": "error", "source": filename, "category": category})
    return pages


def process_excel(file_bytes: bytes, filename: str, category: str) -> list[dict]:
    pages = []
    try:
        ext = Path(filename).suffix.lower()
        df  = (pd.read_csv(BytesIO(file_bytes)) if ext == ".csv"
               else pd.read_excel(BytesIO(file_bytes)))
        df  = df.fillna("")
        # Convert setiap 30 baris ke text
        for i in range(0, len(df), 30):
            chunk = df.iloc[i:i + 30].to_string(index=False)
            if chunk.strip():
                pages.append({"text": chunk, "page": i // 30 + 1,
                               "type": "text", "source": filename, "category": category})
    except Exception as e:
        pages.append({"text": f"Error: {e}", "page": 1,
                      "type": "error", "source": filename, "category": category})
    return pages


def process_image(file_bytes: bytes, filename: str, category: str) -> list[dict]:
    """Proses file gambar (JPG/PNG) dengan Claude OCR."""
    ocr_text = ocr_claude(file_bytes, 1, filename)
    if "[OCR Error" not in ocr_text and ocr_text.strip():
        return [{"text": ocr_text, "page": 1,
                 "type": "claude_ocr", "source": filename, "category": category}]
    return []


# ══════════════════════════════════════════════════════════════
# CHUNKING — CERDAS UNTUK TABEL TEKNIS
# ══════════════════════════════════════════════════════════════

def is_table_line(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    if "|" in s:
        return True
    parts = [p.strip() for p in re.split(r"\s{2,}|\t", s) if p.strip()]
    return len(parts) >= 2


def smart_chunk(text: str, chunk_size: int = CHUNK_SIZE,
                overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Chunking yang mempertahankan integritas tabel parameter."""
    text  = re.sub(r"\n{3,}", "\n\n", text).strip()
    lines = text.split("\n")

    # Deteksi konten tabel
    table_lines = sum(1 for l in lines if is_table_line(l))
    is_table    = table_lines >= 3 or text.count("|") > 5

    if is_table:
        # Pisahkan header dari data rows
        headers   = []
        data_rows = []
        found_data = False
        header_kws = {"parameter", "unit", "normal", "alarm", "interlock",
                       "satuan", "nilai", "gauging", "measuring", "symbol",
                       "clearance", "remark", "limit", "upper", "lower"}

        for ln in lines:
            ls = ln.strip().lower()
            if not found_data and any(kw in ls for kw in header_kws):
                headers.append(ln)
            else:
                found_data = True
                if ln.strip():
                    data_rows.append(ln)

        if not headers and data_rows:
            headers   = data_rows[:2]
            data_rows = data_rows[2:]

        header_str = "\n".join(headers)
        chunks, cur_rows, cur_len = [], [], len(header_str) + 1

        for row in data_rows:
            row_len = len(row) + 1
            if cur_len + row_len > chunk_size and cur_rows:
                content = (header_str + "\n" + "\n".join(cur_rows)).strip()
                if len(content) > 50:
                    chunks.append(content)
                cur_rows = cur_rows[-3:] + [row]
                cur_len  = len(header_str) + sum(len(r) + 1 for r in cur_rows)
            else:
                cur_rows.append(row)
                cur_len += row_len

        if cur_rows:
            content = (header_str + "\n" + "\n".join(cur_rows)).strip()
            if len(content) > 50:
                chunks.append(content)
        return chunks if chunks else [text]

    # Teks biasa — split per kalimat dengan overlap
    sents  = re.split(r"(?<=[.!?\n])\s+", text)
    chunks, cur = [], ""
    for sent in sents:
        if len(cur) + len(sent) + 1 > chunk_size and cur:
            chunks.append(cur.strip())
            words = cur.split()
            cur   = " ".join(words[max(0, len(words) - overlap // 5):]) + " " + sent
        else:
            cur = (cur + " " + sent).strip()
    if cur.strip():
        chunks.append(cur.strip())
    return [c for c in chunks if len(c) > 60]


# ══════════════════════════════════════════════════════════════
# RAG ENGINE
# ══════════════════════════════════════════════════════════════

class RAGEngine:
    def __init__(self, index_dir: Path = INDEX_DIR):
        self.index_dir    = index_dir
        self._embed       = None
        self._faiss       = None
        self._chunks: list[dict] = []
        self._load()

    # ── Embedding model ──────────────────────────────────────
    @property
    def embed(self):
        if self._embed is None:
            try:
                from sentence_transformers import SentenceTransformer
                self._embed = SentenceTransformer(EMBED_MODEL)
            except ImportError:
                pass
        return self._embed

    # ── Persistensi lokal ────────────────────────────────────
    def _save_local(self):
        import faiss
        with open(self.index_dir / "chunks.pkl", "wb") as f:
            pickle.dump(self._chunks, f)
        if self._faiss is not None:
            faiss.write_index(self._faiss, str(self.index_dir / "faiss.index"))

    def _load(self):
        """Load dari lokal, atau pull dari DB jika lokal kosong."""
        # Coba pull dari database jika lokal kosong
        if not (self.index_dir / "chunks.pkl").exists() and is_db_connected():
            db_pull_index(self.index_dir)
        self._load_local()

    def _load_local(self):
        try:
            import faiss
            cp = self.index_dir / "chunks.pkl"
            fp = self.index_dir / "faiss.index"
            if cp.exists():
                with open(cp, "rb") as f:
                    self._chunks = pickle.load(f)
            if fp.exists():
                self._faiss = faiss.read_index(str(fp))
        except Exception:
            self._chunks = []
            self._faiss  = None

    def _save(self):
        """Simpan lokal + push ke database."""
        self._save_local()
        if is_db_connected():
            db_push_index(self.index_dir)

    # ── Embed dan tambahkan ──────────────────────────────────
    def _embed_add(self, new_chunks: list[dict]):
        import faiss
        if not new_chunks or self.embed is None:
            return
        texts  = [c["text"] for c in new_chunks]
        vecs   = self.embed.encode(texts, batch_size=32, show_progress_bar=False)
        vecs   = np.array(vecs, dtype=np.float32)
        norms  = np.linalg.norm(vecs, axis=1, keepdims=True)
        norms[norms == 0] = 1
        vecs  /= norms
        if self._faiss is None:
            self._faiss = faiss.IndexFlatIP(vecs.shape[1])
        self._faiss.add(vecs)
        self._chunks.extend(new_chunks)

    # ── Add dokumen ──────────────────────────────────────────
    def add_document(self, file_bytes: bytes, filename: str, category: str,
                     description: str = "", uploaded_by: str = "",
                     progress_cb=None) -> dict:
        if self.embed is None:
            return {"ok": False, "msg": "Embedding model tidak tersedia"}

        file_hash = hashlib.md5(file_bytes).hexdigest()
        if db_hash_exists(file_hash):
            return {"ok": False, "msg": f"'{filename}' sudah ada di knowledge base"}

        doc_id = "doc_" + file_hash[:12]
        ext    = Path(filename).suffix.lower()

        # Proses sesuai tipe file
        if ext == ".pdf":
            pages = process_pdf(file_bytes, filename, category, progress_cb)
        elif ext in (".docx", ".doc"):
            pages = process_docx(file_bytes, filename, category)
        elif ext in (".xlsx", ".xls", ".csv"):
            pages = process_excel(file_bytes, filename, category)
        elif ext in (".jpg", ".jpeg", ".png"):
            pages = process_image(file_bytes, filename, category)
        elif ext == ".txt":
            text  = file_bytes.decode("utf-8", errors="ignore")
            words = text.split()
            pages = [{"text": " ".join(words[i:i+500]), "page": i//500+1,
                      "type": "text", "source": filename, "category": category}
                     for i in range(0, len(words), 500) if words[i:i+500]]
        else:
            return {"ok": False, "msg": f"Format file tidak didukung: {ext}"}

        if not pages:
            return {"ok": False, "msg": "Tidak ada konten yang bisa diekstrak"}

        # Chunking
        new_chunks = []
        for pg in pages:
            if pg["type"] == "error":
                continue
            for ct in smart_chunk(pg["text"]):
                new_chunks.append({
                    "text":      ct,
                    "source":    filename,
                    "page":      pg["page"],
                    "page_type": pg["type"],
                    "category":  category,
                    "doc_id":    doc_id,
                })

        if not new_chunks:
            return {"ok": False, "msg": "Tidak ada chunk yang terbentuk"}

        self._embed_add(new_chunks)
        self._save()

        n_t = sum(1 for p in pages if p["type"] == "text")
        n_o = sum(1 for p in pages if p["type"] == "claude_ocr")
        n_e = sum(1 for p in pages if p["type"] == "error")

        meta = {
            "name": filename, "category": category, "description": description,
            "pages": len(pages), "n_chunks": len(new_chunks),
            "n_text": n_t, "n_ocr": n_o, "n_error": n_e,
            "hash": file_hash, "ocr_engine": "claude_vision",
            "uploaded_by": uploaded_by,
        }
        db_save_doc_meta(doc_id, meta)

        return {
            "ok": True, "doc_id": doc_id,
            "n_chunks": len(new_chunks), "n_text": n_t,
            "n_ocr": n_o, "n_error": n_e,
            "msg": (f"✅ {len(new_chunks)} chunks dari {len(pages)} halaman "
                    f"({n_t} teks + {n_o} OCR + {n_e} error)"),
        }

    # ── Delete dokumen ───────────────────────────────────────
    def delete_document(self, doc_id: str) -> bool:
        import faiss
        keep = [c for c in self._chunks if c.get("doc_id") != doc_id]
        db_delete_doc(doc_id)
        if not keep:
            self._chunks = []
            self._faiss  = None
            self._save()
            return True
        if self.embed is None:
            return False
        texts  = [c["text"] for c in keep]
        vecs   = self.embed.encode(texts, batch_size=32, show_progress_bar=False)
        vecs   = np.array(vecs, dtype=np.float32)
        norms  = np.linalg.norm(vecs, axis=1, keepdims=True)
        norms[norms == 0] = 1
        vecs  /= norms
        new_idx = faiss.IndexFlatIP(vecs.shape[1])
        new_idx.add(vecs)
        self._chunks = keep
        self._faiss  = new_idx
        self._save()
        return True

    # ── Retrieval ────────────────────────────────────────────
    def retrieve(self, query: str, top_k: int = TOP_K,
                 category: Optional[str] = None) -> list[dict]:
        if not self._chunks or self._faiss is None or self.embed is None:
            return []

        # Query expansion dengan sinonim teknis
        queries = _expand_query(query)
        q_vecs  = self.embed.encode(queries, show_progress_bar=False)
        q_vecs  = np.array(q_vecs, dtype=np.float32)
        norms   = np.linalg.norm(q_vecs, axis=1, keepdims=True)
        norms[norms == 0] = 1
        q_vecs /= norms

        k          = min(top_k * 8, len(self._chunks))
        score_map: dict[int, float] = {}
        for qv in q_vecs:
            scores, idxs = self._faiss.search(qv.reshape(1, -1), k)
            for sc, idx in zip(scores[0], idxs[0]):
                if idx >= 0:
                    score_map[idx] = max(score_map.get(idx, 0.0), float(sc))

        results, seen = [], set()
        for idx, score in sorted(score_map.items(), key=lambda x: x[1], reverse=True):
            if idx >= len(self._chunks):
                continue
            c = self._chunks[idx]
            if category and category != "Semua" and c.get("category", "") != category:
                continue
            key = c["text"][:80]
            if key in seen:
                continue
            seen.add(key)
            results.append({**c, "score": score})
            if len(results) >= top_k:
                break
        return results

    def keyword_search(self, query: str, top_k: int = TOP_K,
                       category: Optional[str] = None) -> list[dict]:
        """Pencarian keyword eksak sebagai fallback."""
        keywords = [w.lower() for w in query.split() if len(w) > 2]
        if not keywords:
            return []
        results = []
        for c in self._chunks:
            if category and category != "Semua" and c.get("category", "") != category:
                continue
            tl    = c["text"].lower()
            score = sum(1 for kw in keywords if kw in tl) / len(keywords)
            if score > 0:
                results.append({**c, "score": score})
        return sorted(results, key=lambda x: x["score"], reverse=True)[:top_k]

    # ── Properties ───────────────────────────────────────────
    @property
    def n_chunks(self):
        return len(self._chunks)

    @property
    def n_ocr_chunks(self):
        return sum(1 for c in self._chunks if c.get("page_type") == "claude_ocr")

    @property
    def categories(self):
        return sorted(set(c.get("category", "") for c in self._chunks if c.get("category")))


# ══════════════════════════════════════════════════════════════
# QUERY EXPANSION — SINONIM TEKNIS
# ══════════════════════════════════════════════════════════════

_SYNONYMS = {
    "thrust pad":    ["推力瓦", "thrust bearing pad", "pad temperature", "瓦温"],
    "bearing":       ["bantalan", "瓦", "bearing temperature", "bearing pad"],
    "lube oil":      ["lubricating oil", "minyak pelumas", "润滑油", "oil pressure"],
    "clearance":     ["celah", "間隙", "tolerance", "fit", "gap"],
    "vibration":     ["vibrasi", "振动", "amplitude", "getaran"],
    "alignment":     ["senter", "对中", "centering", "kopling"],
    "alarm":         ["报警", "high alarm", "low alarm", "batas atas", "batas bawah"],
    "trip":          ["interlock", "跳机", "shutdown", "proteksi", "protection"],
    "temperature":   ["suhu", "temperatur", "temp", "°C", "温度"],
    "pressure":      ["tekanan", "压力", "MPa", "kPa", "bar"],
    "normal":        ["nilai normal", "normal value", "rated", "设计值"],
    "axial":         ["axial displacement", "轴向位移", "thrust position"],
}


def _expand_query(query: str) -> list[str]:
    queries = [query]
    ql = query.lower()
    for key, syns in _SYNONYMS.items():
        if key in ql:
            for syn in syns[:2]:
                exp = re.sub(re.escape(key), syn, ql, flags=re.IGNORECASE)
                if exp not in queries:
                    queries.append(exp)
    # Bersihkan kata tanya
    clean = re.sub(r"\b(berapa|apa|bagaimana|jelaskan|sebutkan|apakah)\b", "", ql).strip()
    if clean and clean not in queries:
        queries.append(clean)
    return queries[:4]


# ══════════════════════════════════════════════════════════════
# QUERY FUNCTION
# ══════════════════════════════════════════════════════════════

def ask_knowledge_base(question: str, engine: RAGEngine,
                       category: Optional[str] = None,
                       context: str = "") -> str:
    import anthropic

    # Semantic search
    chunks = engine.retrieve(question, top_k=TOP_K, category=category)

    # Keyword fallback jika skor rendah
    if not chunks or chunks[0]["score"] < 0.25:
        kw = engine.keyword_search(question, top_k=5, category=category)
        seen = {c["text"][:80] for c in chunks}
        for r in kw:
            if r["text"][:80] not in seen:
                chunks.append(r)
                seen.add(r["text"][:80])
        chunks = chunks[:TOP_K]

    if not chunks:
        return ("⚠️ Tidak ditemukan referensi relevan di knowledge base.\n\n"
                "Pastikan dokumen sudah diupload dan dicek di tab **📑 Dokumen**.")

    # Susun konteks
    ctx_parts = []
    for i, c in enumerate(chunks):
        tag = " [OCR]" if c.get("page_type") == "claude_ocr" else ""
        src = f"[{c['source']} | Hal.{c['page']}{tag} | {c['category']}]"
        ctx_parts.append(f"REFERENSI {i+1} {src}:\n{c['text']}")
    ctx_str = "\n\n---\n\n".join(ctx_parts)

    is_table_req = any(kw in question.lower() for kw in
        ["tunjukkan", "tampilkan", "semua", "seluruh", "daftar",
         "tabel", "list", "all parameter", "batas parameter"])

    system = """Kamu adalah AI Engineer senior spesialis PLTU (Pembangkit Listrik Tenaga Uap).
Keahlian: Boiler CFB, Steam Turbine, Generator, sistem kontrol, maintenance.

ATURAN MENJAWAB:
1. Jawab HANYA berdasarkan REFERENSI yang diberikan.
2. Jika tidak ada di referensi → tulis "Tidak ditemukan di dokumen yang tersedia."
3. Sebutkan sumber (nama dokumen + halaman) di akhir jawaban.
4. Satuan teknis WAJIB akurat: MPa, kPa, °C, bar, rpm, mm, kN·m, MW.

FORMAT TABEL PARAMETER (gunakan selalu untuk data parameter):
| No | Parameter | Satuan | Normal | Alarm High | Alarm Low | Trip | Keterangan |
|----|-----------|--------|--------|------------|-----------|------|------------|

ATURAN TABEL:
- Tampilkan SEMUA parameter dari referensi, jangan ada yang terlewat.
- Nilai "-" jika tidak disebutkan di referensi.
- Urutkan: Temperatur → Tekanan → Level → Flow → Lainnya."""

    user_msg = f"PERTANYAAN: {question}\n\n"
    if context:
        user_msg += f"DATA OPERASIONAL SAAT INI:\n{context}\n\n"
    if is_table_req:
        user_msg += "INSTRUKSI: Tampilkan dalam format tabel lengkap, kumpulkan dari SEMUA referensi.\n\n"
    user_msg += f"REFERENSI DOKUMEN:\n{ctx_str}"

    try:
        client   = anthropic.Anthropic(api_key=st.secrets["anthropic"]["api_key"])
        max_tok  = 4000 if is_table_req else 2000
        response = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=max_tok,
            system=system,
            messages=[{"role": "user", "content": user_msg}],
        )
        answer  = response.content[0].text
        sources = list(dict.fromkeys(f"{c['source']} (Hal.{c['page']})" for c in chunks))
        return answer + f"\n\n---\n📚 **Sumber:** {' · '.join(sources[:6])}"
    except Exception as e:
        return f"❌ Error API: {e}"


# ══════════════════════════════════════════════════════════════
# STREAMLIT UI
# ══════════════════════════════════════════════════════════════

# ── Load engine (cached) ───────────────────────────────────
@st.cache_resource(show_spinner=False)
def load_engine() -> RAGEngine:
    return RAGEngine()



# ══════════════════════════════════════════════════════════════
# PROFESSIONAL UI
# ══════════════════════════════════════════════════════════════

THEME_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

/* ── Reset & Base ── */
* { font-family: 'Inter', sans-serif !important; box-sizing: border-box; }
.stApp { background: #0a0e1a !important; }

/* ── Hide Streamlit chrome ── */
#MainMenu, footer, header { visibility: hidden; }
.stDeployButton { display: none; }

/* ── Sidebar ── */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0d1117 0%, #161b22 100%) !important;
    border-right: 1px solid #21262d !important;
    width: 280px !important;
}
section[data-testid="stSidebar"] > div { padding: 0 !important; }

/* ── Main content padding ── */
.block-container { padding: 2rem 2.5rem 2rem 2.5rem !important; max-width: 1400px; }

/* ── Typography ── */
h1 { font-size: 1.75rem !important; font-weight: 800 !important; color: #f0f6fc !important; letter-spacing: -0.5px; }
h2 { font-size: 1.2rem !important; font-weight: 700 !important; color: #e6edf3 !important; }
h3 { font-size: 1rem !important; font-weight: 600 !important; color: #8b949e !important; text-transform: uppercase; letter-spacing: 0.5px; }
p, .stMarkdown { color: #c9d1d9 !important; }

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
    background: #161b22 !important;
    border-radius: 12px !important;
    padding: 4px !important;
    gap: 2px !important;
    border: 1px solid #21262d !important;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 9px !important;
    font-weight: 600 !important;
    font-size: 0.82rem !important;
    color: #8b949e !important;
    padding: 8px 16px !important;
    border: none !important;
    transition: all 0.2s !important;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #238636, #2ea043) !important;
    color: #ffffff !important;
    box-shadow: 0 2px 8px rgba(35,134,54,0.4) !important;
}

/* ── Buttons ── */
.stButton > button {
    background: linear-gradient(135deg, #238636, #2ea043) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    padding: 8px 20px !important;
    transition: all 0.2s !important;
    box-shadow: 0 2px 8px rgba(35,134,54,0.3) !important;
}
.stButton > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 12px rgba(35,134,54,0.5) !important;
}
.stButton > button[kind="secondary"] {
    background: #21262d !important;
    color: #c9d1d9 !important;
    border: 1px solid #30363d !important;
    box-shadow: none !important;
}

/* ── Inputs ── */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    background: #161b22 !important;
    color: #e6edf3 !important;
    border: 1px solid #30363d !important;
    border-radius: 8px !important;
    font-size: 0.9rem !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: #388bfd !important;
    box-shadow: 0 0 0 3px rgba(56,139,253,0.15) !important;
}
.stSelectbox > div > div,
.stMultiSelect > div > div {
    background: #161b22 !important;
    border: 1px solid #30363d !important;
    border-radius: 8px !important;
    color: #e6edf3 !important;
}

/* ── File uploader ── */
[data-testid="stFileUploaderDropzone"] {
    background: #161b22 !important;
    border: 2px dashed #30363d !important;
    border-radius: 12px !important;
    transition: all 0.2s !important;
}
[data-testid="stFileUploaderDropzone"]:hover {
    border-color: #388bfd !important;
    background: #0d1117 !important;
}

/* ── Metrics ── */
[data-testid="metric-container"] {
    background: #161b22 !important;
    border: 1px solid #21262d !important;
    border-radius: 10px !important;
    padding: 14px 16px !important;
}
[data-testid="metric-container"] label {
    color: #8b949e !important;
    font-size: 0.72rem !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.8px !important;
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
    color: #58a6ff !important;
    font-size: 1.6rem !important;
    font-weight: 800 !important;
}

/* ── Alerts ── */
.stSuccess { background: rgba(35,134,54,0.1) !important; border: 1px solid rgba(35,134,54,0.3) !important; border-radius: 8px !important; }
.stWarning { background: rgba(210,153,34,0.1) !important; border: 1px solid rgba(210,153,34,0.3) !important; border-radius: 8px !important; }
.stError   { background: rgba(248,81,73,0.1) !important; border: 1px solid rgba(248,81,73,0.3) !important; border-radius: 8px !important; }
.stInfo    { background: rgba(56,139,253,0.1) !important; border: 1px solid rgba(56,139,253,0.3) !important; border-radius: 8px !important; }

/* ── Chat messages ── */
[data-testid="stChatMessage"] {
    background: #161b22 !important;
    border: 1px solid #21262d !important;
    border-radius: 12px !important;
    margin-bottom: 8px !important;
}
[data-testid="stChatInput"] {
    background: #161b22 !important;
    border: 1px solid #30363d !important;
    border-radius: 12px !important;
}

/* ── Expander ── */
.streamlit-expanderHeader {
    background: #161b22 !important;
    border: 1px solid #21262d !important;
    border-radius: 8px !important;
    color: #e6edf3 !important;
    font-weight: 600 !important;
}

/* ── Divider ── */
hr { border-color: #21262d !important; margin: 1.5rem 0 !important; }

/* ── Progress ── */
.stProgress > div > div { background: linear-gradient(90deg, #238636, #2ea043) !important; }

/* ── Scrollbar ── */
::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: #0d1117; }
::-webkit-scrollbar-thumb { background: #30363d; border-radius: 3px; }
::-webkit-scrollbar-thumb:hover { background: #484f58; }

/* ── Code blocks ── */
.stCodeBlock { background: #161b22 !important; border: 1px solid #30363d !important; border-radius: 8px !important; }

/* ── Table ── */
.stDataFrame { background: #161b22 !important; border: 1px solid #21262d !important; border-radius: 8px !important; }
</style>
"""

# ── Custom HTML components ─────────────────────────────────────

def _sidebar_header(n_docs, n_chunks, n_ocr, db_ok, username):
    db_color  = "#2ea043" if db_ok else "#f85149"
    db_label  = "Connected" if db_ok else "Disconnected"
    db_dot    = "🟢" if db_ok else "🔴"
    return f"""
<div style="padding: 20px 16px 0;">
    <div style="display:flex;align-items:center;gap:10px;margin-bottom:4px;">
        <div style="width:32px;height:32px;background:linear-gradient(135deg,#238636,#2ea043);
            border-radius:8px;display:flex;align-items:center;justify-content:center;
            font-size:16px;">📚</div>
        <div>
            <div style="font-size:0.95rem;font-weight:700;color:#f0f6fc;">DIGIT-OPS</div>
            <div style="font-size:0.7rem;color:#8b949e;font-weight:500;">Knowledge Base</div>
        </div>
    </div>
    <div style="margin:14px 0;padding:8px 12px;background:rgba({('35,134,54' if db_ok else '248,81,73')},0.1);
        border:1px solid rgba({('35,134,54' if db_ok else '248,81,73')},0.3);
        border-radius:8px;display:flex;align-items:center;gap:8px;">
        <span style="font-size:0.75rem;color:{'#2ea043' if db_ok else '#f85149'};font-weight:600;">
            {db_dot} Database {db_label}
        </span>
    </div>
    <div style="display:grid;grid-template-columns:1fr 1fr;gap:8px;margin-bottom:4px;">
        <div style="background:#161b22;border:1px solid #21262d;border-radius:8px;
            padding:10px 12px;text-align:center;">
            <div style="font-size:1.4rem;font-weight:800;color:#58a6ff;">{n_docs}</div>
            <div style="font-size:0.65rem;color:#8b949e;font-weight:600;text-transform:uppercase;letter-spacing:0.8px;">Dokumen</div>
        </div>
        <div style="background:#161b22;border:1px solid #21262d;border-radius:8px;
            padding:10px 12px;text-align:center;">
            <div style="font-size:1.4rem;font-weight:800;color:#58a6ff;">{n_chunks:,}</div>
            <div style="font-size:0.65rem;color:#8b949e;font-weight:600;text-transform:uppercase;letter-spacing:0.8px;">Chunks</div>
        </div>
        <div style="background:#161b22;border:1px solid #21262d;border-radius:8px;
            padding:10px 12px;text-align:center;">
            <div style="font-size:1.4rem;font-weight:800;color:#d29922;">{n_ocr:,}</div>
            <div style="font-size:0.65rem;color:#8b949e;font-weight:600;text-transform:uppercase;letter-spacing:0.8px;">OCR Pages</div>
        </div>
        <div style="background:#161b22;border:1px solid #21262d;border-radius:8px;
            padding:10px 12px;text-align:center;">
            <div style="font-size:1.4rem;font-weight:800;color:#3fb950;">{n_chunks - n_ocr:,}</div>
            <div style="font-size:0.65rem;color:#8b949e;font-weight:600;text-transform:uppercase;letter-spacing:0.8px;">Text Pages</div>
        </div>
    </div>
</div>
"""


def _doc_card(doc: dict) -> str:
    name     = doc.get("name", "—")
    cat      = doc.get("category", "—")
    pages    = doc.get("pages", 0)
    chunks   = doc.get("n_chunks", 0)
    n_ocr    = doc.get("n_ocr", 0)
    n_text   = doc.get("n_text", 0)
    desc     = doc.get("description", "") or ""
    uploader = doc.get("uploaded_by", "") or "—"
    date_raw = doc.get("created_at", "")
    date_str = str(date_raw)[:10] if date_raw else "—"
    pct_ocr  = int(n_ocr / max(pages, 1) * 100)

    ext = Path(name).suffix.lower()
    icon_map = {".pdf": "📄", ".docx": "📝", ".doc": "📝",
                ".xlsx": "📊", ".xls": "📊", ".csv": "📊",
                ".jpg": "🖼️", ".jpeg": "🖼️", ".png": "🖼️", ".txt": "📃"}
    icon = icon_map.get(ext, "📁")

    return f"""
<div style="background:#161b22;border:1px solid #21262d;border-radius:12px;
     padding:16px 18px;margin-bottom:10px;transition:border-color 0.2s;"
     onmouseover="this.style.borderColor='#388bfd'"
     onmouseout="this.style.borderColor='#21262d'">
    <div style="display:flex;align-items:flex-start;justify-content:space-between;gap:12px;">
        <div style="flex:1;min-width:0;">
            <div style="display:flex;align-items:center;gap:8px;margin-bottom:6px;">
                <span style="font-size:1.1rem;">{icon}</span>
                <span style="font-size:0.9rem;font-weight:600;color:#e6edf3;
                    overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">{name}</span>
            </div>
            <div style="display:flex;flex-wrap:wrap;gap:6px;margin-bottom:8px;">
                <span style="background:rgba(56,139,253,0.12);color:#388bfd;
                    border:1px solid rgba(56,139,253,0.25);border-radius:20px;
                    padding:2px 10px;font-size:0.72rem;font-weight:500;">{cat}</span>
                <span style="background:#21262d;color:#8b949e;border-radius:20px;
                    padding:2px 10px;font-size:0.72rem;">📃 {pages} hal</span>
                <span style="background:#21262d;color:#8b949e;border-radius:20px;
                    padding:2px 10px;font-size:0.72rem;">🧩 {chunks} chunks</span>
                <span style="background:rgba(210,153,34,0.12);color:#d29922;border-radius:20px;
                    padding:2px 10px;font-size:0.72rem;">🔮 {n_ocr} OCR</span>
            </div>
            {f'<div style="font-size:0.78rem;color:#8b949e;font-style:italic;margin-bottom:6px;">{desc}</div>' if desc else ''}
            <div style="display:flex;gap:12px;">
                <span style="font-size:0.72rem;color:#6e7681;">👤 {uploader}</span>
                <span style="font-size:0.72rem;color:#6e7681;">📅 {date_str}</span>
            </div>
        </div>
    </div>
    <div style="margin-top:10px;">
        <div style="display:flex;justify-content:space-between;margin-bottom:4px;">
            <span style="font-size:0.7rem;color:#6e7681;">OCR Coverage</span>
            <span style="font-size:0.7rem;color:#6e7681;">{pct_ocr}%</span>
        </div>
        <div style="background:#21262d;border-radius:4px;height:4px;">
            <div style="background:linear-gradient(90deg,#d29922,#e3b341);
                height:4px;border-radius:4px;width:{pct_ocr}%;transition:width 0.3s;"></div>
        </div>
    </div>
</div>"""


def _search_result_card(r: dict, idx: int) -> str:
    pct    = min(100, int(r["score"] * 100))
    pt     = r.get("page_type", "text")
    is_ocr = pt == "claude_ocr"
    tag_bg  = "rgba(210,153,34,0.12)" if is_ocr else "rgba(46,160,67,0.12)"
    tag_col = "#d29922"               if is_ocr else "#2ea043"
    tag_txt = "🔮 Claude OCR"         if is_ocr else "📝 Text"

    bar_col = ("#2ea043" if pct > 70 else "#d29922" if pct > 45 else "#f85149")
    snippet = r["text"][:350].replace("<", "&lt;").replace(">", "&gt;")
    if len(r["text"]) > 350:
        snippet += "…"

    return f"""
<div style="background:#161b22;border:1px solid #21262d;border-radius:10px;
     padding:14px 16px;margin-bottom:8px;border-left:3px solid {bar_col};">
    <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:8px;">
        <div style="display:flex;align-items:center;gap:8px;flex-wrap:wrap;">
            <span style="font-size:0.8rem;font-weight:600;color:#e6edf3;">#{idx+1}</span>
            <span style="font-size:0.78rem;color:#58a6ff;font-weight:500;">{r['source']}</span>
            <span style="background:#21262d;color:#8b949e;border-radius:12px;
                padding:2px 8px;font-size:0.7rem;">Hal. {r['page']}</span>
            <span style="background:{tag_bg};color:{tag_col};border-radius:12px;
                padding:2px 8px;font-size:0.7rem;">{tag_txt}</span>
            <span style="background:#21262d;color:#8b949e;border-radius:12px;
                padding:2px 8px;font-size:0.7rem;">{r.get('category','')[:30]}</span>
        </div>
        <div style="display:flex;align-items:center;gap:6px;flex-shrink:0;">
            <div style="background:#21262d;border-radius:4px;height:6px;width:60px;">
                <div style="background:{bar_col};height:6px;border-radius:4px;width:{pct}%;"></div>
            </div>
            <span style="font-size:0.78rem;font-weight:700;color:{bar_col};">{pct}%</span>
        </div>
    </div>
    <div style="font-size:0.82rem;color:#8b949e;line-height:1.6;
        font-family:ui-monospace,monospace;white-space:pre-wrap;">{snippet}</div>
</div>"""


def _section_header(icon: str, title: str, subtitle: str = "") -> str:
    return f"""
<div style="margin-bottom:20px;">
    <div style="display:flex;align-items:center;gap:10px;">
        <div style="width:36px;height:36px;background:rgba(35,134,54,0.15);
            border:1px solid rgba(35,134,54,0.3);border-radius:8px;
            display:flex;align-items:center;justify-content:center;font-size:16px;">{icon}</div>
        <div>
            <h2 style="margin:0;font-size:1.15rem;font-weight:700;color:#e6edf3;">{title}</h2>
            {f'<p style="margin:0;font-size:0.78rem;color:#8b949e;">{subtitle}</p>' if subtitle else ''}
        </div>
    </div>
</div>"""


def _shortcut_card(icon, title, subtitle, color):
    return f"""
<div style="background:#161b22;border:1px solid #21262d;border-radius:10px;
     padding:14px 16px;margin-bottom:8px;cursor:pointer;
     border-left:3px solid {color};">
    <div style="font-size:1.1rem;margin-bottom:4px;">{icon}</div>
    <div style="font-size:0.85rem;font-weight:600;color:#e6edf3;">{title}</div>
    <div style="font-size:0.74rem;color:#8b949e;margin-top:2px;">{subtitle}</div>
</div>"""


# ── Main app ────────────────────────────────────────────────────

@st.cache_resource(show_spinner=False)
def load_engine() -> "RAGEngine":
    return RAGEngine()


def main():
    import streamlit as st

    st.markdown(THEME_CSS, unsafe_allow_html=True)

    engine = load_engine()
    db_ok  = is_db_connected()
    docs   = db_get_all_docs()

    # ════════════════════════════════════════════════════════
    # SIDEBAR
    # ════════════════════════════════════════════════════════
    with st.sidebar:
        uname = st.session_state.get("username", "")

        st.markdown(
            _sidebar_header(len(docs), engine.n_chunks,
                            engine.n_ocr_chunks, db_ok, uname),
            unsafe_allow_html=True,
        )

        st.markdown("""
<div style="padding:0 16px;">
<hr style="border-color:#21262d;margin:12px 0;">
</div>""", unsafe_allow_html=True)

        with st.container():
            st.markdown(
                '<div style="padding:0 16px;"><p style="font-size:0.72rem;'
                'color:#8b949e;font-weight:600;text-transform:uppercase;'
                'letter-spacing:0.8px;margin-bottom:8px;">Filter Kategori</p></div>',
                unsafe_allow_html=True,
            )
            cat_opts  = ["Semua"] + engine.categories
            sel_cat   = st.selectbox("", cat_opts, key="sel_cat",
                                      label_visibility="collapsed")

        st.markdown("""
<div style="padding:0 16px;">
<hr style="border-color:#21262d;margin:12px 0;">
<p style="font-size:0.72rem;color:#8b949e;font-weight:600;text-transform:uppercase;
letter-spacing:0.8px;margin-bottom:8px;">Sinkronisasi Database</p>
</div>""", unsafe_allow_html=True)

        with st.container():
            col_p, col_u = st.columns(2)
            if col_p.button("⬆ Push", key="btn_push", use_container_width=True):
                with st.spinner(""):
                    ok, msg = db_push_index(INDEX_DIR)
                st.toast(msg)
            if col_u.button("⬇ Pull", key="btn_pull", use_container_width=True):
                with st.spinner(""):
                    ok, msg = db_pull_index(INDEX_DIR)
                st.toast(msg)
                if ok:
                    st.cache_resource.clear()
                    st.rerun()

        st.markdown("""
<div style="padding:0 16px;">
<hr style="border-color:#21262d;margin:12px 0;">
<p style="font-size:0.72rem;color:#8b949e;font-weight:600;text-transform:uppercase;
letter-spacing:0.8px;margin-bottom:8px;">Profil</p>
</div>""", unsafe_allow_html=True)

        with st.container():
            st.text_input("", key="username", placeholder="Nama Anda...",
                          label_visibility="collapsed")

        st.markdown("""
<div style="padding:16px;position:absolute;bottom:0;left:0;right:0;">
<div style="font-size:0.68rem;color:#6e7681;text-align:center;">
DIGIT-OPS Knowledge Base v2.0<br>
PLN Nusantara Power Services
</div>
</div>""", unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════
    # HEADER
    # ════════════════════════════════════════════════════════
    st.markdown("""
<div style="margin-bottom:24px;">
    <div style="display:flex;align-items:center;justify-content:space-between;
        flex-wrap:wrap;gap:12px;">
        <div>
            <h1 style="margin:0;font-size:1.6rem;font-weight:800;color:#f0f6fc;
                background:linear-gradient(135deg,#58a6ff,#79c0ff);
                -webkit-background-clip:text;-webkit-text-fill-color:transparent;">
                📚 DIGIT-OPS Knowledge Base
            </h1>
            <p style="margin:4px 0 0;font-size:0.82rem;color:#8b949e;">
                Manual Book · Jurnal · Referensi Teknis PLTU &nbsp;·&nbsp; Bilingual ID + EN &nbsp;·&nbsp; Claude Vision OCR
            </p>
        </div>
    </div>
</div>""", unsafe_allow_html=True)

    if not db_ok:
        st.warning(
            "⚠️ **Database belum terhubung.** "
            "Tambahkan `DATABASE_URL` di Streamlit Cloud → Settings → Secrets."
        )

    # ════════════════════════════════════════════════════════
    # TABS
    # ════════════════════════════════════════════════════════
    tab_chat, tab_upload, tab_search, tab_docs, tab_setup = st.tabs([
        "💬  Tanya AI",
        "📤  Upload",
        "🔍  Pencarian",
        "📑  Dokumen",
        "⚙️  Setup",
    ])

    # ────────────────────────────────────────────────────────
    # TAB 1 — TANYA AI
    # ────────────────────────────────────────────────────────
    with tab_chat:
        if engine.n_chunks == 0:
            st.markdown("""
<div style="background:rgba(56,139,253,0.08);border:1px solid rgba(56,139,253,0.25);
    border-radius:12px;padding:20px 24px;text-align:center;margin-bottom:20px;">
    <div style="font-size:2rem;margin-bottom:8px;">📭</div>
    <div style="font-size:0.95rem;font-weight:600;color:#58a6ff;margin-bottom:4px;">
        Knowledge Base Kosong
    </div>
    <div style="font-size:0.82rem;color:#8b949e;">
        Upload manual book, jurnal, atau referensi teknis di tab <b>📤 Upload</b> terlebih dahulu.
    </div>
</div>""", unsafe_allow_html=True)

        # Shortcut buttons
        st.markdown(_section_header("⚡", "Shortcut Parameter",
            "Tampilkan semua nilai normal, alarm, dan trip sekaligus"), unsafe_allow_html=True)

        _Q = {
            "boiler": ("Tampilkan SEMUA batas parameter operasi Boiler CFB dalam tabel lengkap "
                       "(nilai normal, alarm high, alarm low, trip, satuan): bed temperature, "
                       "furnace pressure, steam pressure, steam temperature, drum level, "
                       "ID/FD/PA fan, coal feeder, flue gas, O2, NOx."),
            "turbin": ("Tampilkan SEMUA batas parameter operasi Steam Turbine dalam tabel lengkap "
                       "(nilai normal, alarm high, alarm low, trip, satuan): thrust pad temp, "
                       "bearing temp, lube oil pressure/temp, axial displacement, vibration, "
                       "steam pressure/temp, exhaust pressure, speed, scavenge oil temp."),
            "gen":    ("Tampilkan SEMUA batas parameter operasi Generator dalam tabel lengkap "
                       "(nilai normal, alarm high, alarm low, trip, satuan)."),
        }

        sc1, sc2, sc3 = st.columns(3)
        run_q = None
        run_label = None

        with sc1:
            st.markdown(_shortcut_card("🔥", "Parameter Boiler CFB",
                "Bed temp · Pressure · Steam · Fan · Drum", "#f85149"), unsafe_allow_html=True)
            if st.button("Tampilkan →", key="q_boiler", use_container_width=True):
                run_q, run_label = _Q["boiler"], "📊 Semua Parameter Boiler CFB"
        with sc2:
            st.markdown(_shortcut_card("⚙️", "Parameter Steam Turbine",
                "Thrust pad · Bearing · Lube oil · Vibration", "#58a6ff"), unsafe_allow_html=True)
            if st.button("Tampilkan →", key="q_turbin", use_container_width=True):
                run_q, run_label = _Q["turbin"], "📊 Semua Parameter Steam Turbine"
        with sc3:
            st.markdown(_shortcut_card("⚡", "Parameter Generator",
                "Voltage · Current · Temperature · Excitation", "#d29922"), unsafe_allow_html=True)
            if st.button("Tampilkan →", key="q_gen", use_container_width=True):
                run_q, run_label = _Q["gen"], "📊 Semua Parameter Generator"

        if run_q:
            with st.spinner("Mengumpulkan data dari knowledge base..."):
                cat = None if sel_cat == "Semua" else sel_cat
                ans = ask_knowledge_base(run_q, engine, category=cat)
            if "messages" not in st.session_state:
                st.session_state.messages = []
            st.session_state.messages += [
                {"role": "user",      "content": run_label},
                {"role": "assistant", "content": ans},
            ]
            st.rerun()

        st.divider()

        # Chat area
        st.markdown(_section_header("💬", "Tanya Manual Book",
            "Tanyakan dalam bahasa Indonesia atau Inggris"), unsafe_allow_html=True)

        if "messages" not in st.session_state:
            st.session_state.messages = []

        # Render chat history
        chat_container = st.container()
        with chat_container:
            for msg in st.session_state.messages:
                with st.chat_message(msg["role"]):
                    st.markdown(msg["content"])

        # Input
        if prompt := st.chat_input(
                "Contoh: Berapa thrust pad temperature normal dan alarm?"):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            with st.chat_message("assistant"):
                placeholder = st.empty()
                with st.spinner("Mencari referensi..."):
                    cat = None if sel_cat == "Semua" else sel_cat
                    try:
                        ans = ask_knowledge_base(prompt, engine, category=cat)
                    except Exception as e:
                        ans = f"❌ Error: {e}"
                placeholder.markdown(ans)
            st.session_state.messages.append({"role": "assistant", "content": ans})

        # Toolbar
        if st.session_state.messages:
            col_clr, col_exp, _ = st.columns([1, 1, 3])
            if col_clr.button("🗑 Hapus Chat", key="clr", use_container_width=True):
                st.session_state.messages = []
                st.rerun()
            chat_txt = "\n\n".join(
                f"[{'USER' if m['role']=='user' else 'AI'}]\n{m['content']}"
                for m in st.session_state.messages
            )
            col_exp.download_button(
                "💾 Export", data=chat_txt,
                file_name=f"chat_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain", key="dl_chat", use_container_width=True,
            )

    # ────────────────────────────────────────────────────────
    # TAB 2 — UPLOAD
    # ────────────────────────────────────────────────────────
    with tab_upload:
        st.markdown(_section_header("📤", "Upload Dokumen",
            "PDF · DOCX · XLSX · CSV · TXT · JPG · PNG — maks 200MB per file"),
            unsafe_allow_html=True)

        # API check
        try:
            api_ok = bool(st.secrets["anthropic"]["api_key"])
        except Exception:
            api_ok = False

        if not api_ok:
            st.error("❌ **Anthropic API key tidak ditemukan.** "
                     "Tambahkan `[anthropic] api_key` di Secrets.")
            st.stop()

        # Info cards
        ic1, ic2, ic3 = st.columns(3)
        ic1.markdown("""
<div style="background:#161b22;border:1px solid #21262d;border-radius:10px;padding:14px;">
    <div style="font-size:1.2rem;margin-bottom:6px;">📄 PDF</div>
    <div style="font-size:0.78rem;color:#8b949e;line-height:1.5;">
        Halaman teks → ekstrak langsung.<br>
        Halaman gambar/tabel → Claude Vision OCR.
    </div>
</div>""", unsafe_allow_html=True)
        ic2.markdown("""
<div style="background:#161b22;border:1px solid #21262d;border-radius:10px;padding:14px;">
    <div style="font-size:1.2rem;margin-bottom:6px;">📝 DOCX / XLSX</div>
    <div style="font-size:0.78rem;color:#8b949e;line-height:1.5;">
        Word & Excel diekstrak termasuk tabel dan data terstruktur.
    </div>
</div>""", unsafe_allow_html=True)
        ic3.markdown("""
<div style="background:#161b22;border:1px solid #21262d;border-radius:10px;padding:14px;">
    <div style="font-size:1.2rem;margin-bottom:6px;">🖼️ Gambar</div>
    <div style="font-size:0.78rem;color:#8b949e;line-height:1.5;">
        JPG/PNG diproses Claude Vision OCR — ideal untuk foto manual book.
    </div>
</div>""", unsafe_allow_html=True)

        st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)

        # Upload form
        uploaded_files = st.file_uploader(
            "Drag & drop file di sini, atau klik Browse files:",
            accept_multiple_files=True,
            type=["pdf","docx","doc","xlsx","xls","csv","txt","jpg","jpeg","png"],
            key="file_uploader",
        )

        col_cat, col_desc = st.columns([1, 2])
        doc_cat  = col_cat.selectbox("Kategori:", DOC_CATEGORIES, key="doc_cat")
        doc_desc = col_desc.text_input("Deskripsi:", key="doc_desc",
                                        placeholder="contoh: Steam Turbine N27.5 OEM Manual 2019")

        if uploaded_files:
            st.markdown(f"<div style='color:#8b949e;font-size:0.8rem;margin:8px 0;'>"
                        f"✅ {len(uploaded_files)} file siap diproses</div>",
                        unsafe_allow_html=True)

            if st.button(f"📥  Proses & Simpan  ({len(uploaded_files)} file)",
                          key="btn_upload", use_container_width=True, type="primary"):
                uname = st.session_state.get("username", "")
                all_ok = True
                for uf in uploaded_files:
                    fb  = uf.read()
                    ext = Path(uf.name).suffix.lower()

                    st.markdown(f"""
<div style="background:#161b22;border:1px solid #21262d;border-radius:10px;
     padding:14px 18px;margin:8px 0;">
    <div style="font-size:0.88rem;font-weight:600;color:#e6edf3;
        margin-bottom:10px;">📄 {uf.name}</div>""",
                        unsafe_allow_html=True)

                    prog    = st.progress(0.0)
                    cap_ph  = st.empty()
                    log_ph  = st.empty()
                    log_buf = []

                    def _cb(pg, tot, nt, nv,
                            _p=prog, _c=cap_ph, _l=log_ph, _b=log_buf, _fn=uf.name):
                        _p.progress(pg / tot)
                        _c.markdown(
                            f"<div style='font-size:0.78rem;color:#8b949e;'>"
                            f"Halaman <b style='color:#e6edf3;'>{pg}/{tot}</b>"
                            f" &nbsp;·&nbsp; Teks: {nt} &nbsp;·&nbsp; OCR: {nv}</div>",
                            unsafe_allow_html=True,
                        )
                        if nv > 0:
                            _b.append(f"  ✓ Hal {pg}: Claude Vision OCR selesai")
                            _l.code("\n".join(_b[-4:]), language=None)

                    with st.spinner(""):
                        res = engine.add_document(
                            fb, uf.name, doc_cat,
                            description=doc_desc,
                            uploaded_by=uname,
                            progress_cb=(_cb if ext == ".pdf" else None),
                        )

                    prog.empty(); cap_ph.empty(); log_ph.empty()

                    if res["ok"]:
                        m1, m2, m3, m4 = st.columns(4)
                        m1.metric("Chunks",   res["n_chunks"])
                        m2.metric("Hal Teks", res["n_text"])
                        m3.metric("Hal OCR",  res["n_ocr"])
                        m4.metric("Error",    res["n_error"])
                        st.success(res["msg"])
                    else:
                        st.warning(res["msg"])
                        all_ok = False

                    st.markdown("</div>", unsafe_allow_html=True)

                if all_ok:
                    st.cache_resource.clear()
                    st.rerun()

    # ────────────────────────────────────────────────────────
    # TAB 3 — PENCARIAN
    # ────────────────────────────────────────────────────────
    with tab_search:
        st.markdown(_section_header("🔍", "Pencarian Dokumen",
            "Semantik (AI) atau Keyword Eksak"), unsafe_allow_html=True)

        if engine.n_chunks == 0:
            st.info("Upload dokumen terlebih dahulu.")
        else:
            # Search bar
            query = st.text_input(
                "",
                key="search_q",
                placeholder="🔍  Cari parameter, prosedur, nilai teknis...",
                label_visibility="collapsed",
            )

            col_m, col_c, col_n, col_s = st.columns([2, 2, 1, 1])
            mode     = col_m.radio("Mode:", ["🤖 Semantik (AI)", "🔤 Keyword Eksak"],
                                    key="smode", horizontal=True)
            cat_filt = col_c.selectbox("Kategori:", ["Semua"] + engine.categories,
                                        key="scat")
            top_n    = col_n.slider("Hasil:", 3, 20, 8, key="top_n")
            show_txt = col_s.checkbox("Tampilkan teks", key="show_txt", value=True)

            if st.button("🔍  Cari Sekarang", key="btn_search",
                          use_container_width=True, type="primary") and query:
                cat = None if cat_filt == "Semua" else cat_filt
                with st.spinner("Mencari..."):
                    if "Keyword" in mode:
                        res = engine.keyword_search(query, top_k=int(top_n), category=cat)
                    else:
                        res = engine.retrieve(query, top_k=int(top_n), category=cat)
                st.session_state["s_results"] = res
                st.session_state["s_query"]   = query

            results = st.session_state.get("s_results", [])
            if results:
                q_lbl = st.session_state.get("s_query", "")
                st.markdown(
                    f"<div style='font-size:0.8rem;color:#8b949e;margin:12px 0 8px;'>"
                    f"<b style='color:#e6edf3;'>{len(results)}</b> hasil untuk: "
                    f"<i style='color:#58a6ff;'>'{q_lbl}'</i></div>",
                    unsafe_allow_html=True,
                )

                for i, r in enumerate(results):
                    if show_txt:
                        st.markdown(_search_result_card(r, i), unsafe_allow_html=True)
                    else:
                        pct = min(100, int(r["score"] * 100))
                        bar = "🟢" if pct > 70 else "🟡" if pct > 45 else "🔴"
                        st.markdown(
                            f"{bar} **{r['source']}** Hal.{r['page']} "
                            f"— {r.get('category','')[:25]} — **{pct}%**"
                        )

                # Export
                st.divider()
                df_exp = pd.DataFrame([{
                    "Dokumen":   r["source"], "Halaman": r["page"],
                    "Kategori":  r["category"], "Tipe": r.get("page_type","text"),
                    "Relevansi": f"{min(100,int(r['score']*100))}%",
                    "Teks":      r["text"][:300],
                } for r in results])
                st.download_button(
                    "📥 Export Hasil (.csv)", data=df_exp.to_csv(index=False),
                    file_name="hasil_pencarian.csv", mime="text/csv", key="dl_s",
                )

    # ────────────────────────────────────────────────────────
    # TAB 4 — DOKUMEN
    # ────────────────────────────────────────────────────────
    with tab_docs:
        st.markdown(_section_header("📑", "Manajemen Dokumen",
            f"{len(docs)} dokumen tersimpan di knowledge base"), unsafe_allow_html=True)

        if not docs:
            st.info("Belum ada dokumen. Upload di tab **📤 Upload**.")
        else:
            # Summary metrics
            total_pg  = sum(d.get("pages",0)  for d in docs)
            total_ck  = engine.n_chunks
            total_ocr = sum(d.get("n_ocr",0)  for d in docs)
            cats      = len(set(d.get("category","") for d in docs))

            m1, m2, m3, m4, m5 = st.columns(5)
            m1.metric("Dokumen",   len(docs))
            m2.metric("Halaman",   total_pg)
            m3.metric("Chunks",    f"{total_ck:,}")
            m4.metric("OCR Pages", total_ocr)
            m5.metric("Kategori",  cats)

            st.divider()

            # Search & filter
            col_sf, col_cf = st.columns([2, 1])
            doc_q  = col_sf.text_input("", placeholder="🔎 Cari nama / deskripsi...",
                                        key="dq", label_visibility="collapsed")
            cat_df = col_cf.selectbox("", ["Semua"] + sorted(
                                       set(d.get("category","") for d in docs)),
                                       key="dcf", label_visibility="collapsed")

            filtered = [d for d in docs
                        if (not doc_q or
                            doc_q.lower() in d.get("name","").lower() or
                            doc_q.lower() in d.get("description","").lower())
                        and (cat_df == "Semua" or d.get("category") == cat_df)]

            st.markdown(
                f"<div style='font-size:0.78rem;color:#8b949e;margin:8px 0;'>"
                f"Menampilkan <b style='color:#e6edf3;'>{len(filtered)}</b> dokumen</div>",
                unsafe_allow_html=True,
            )

            # List dokumen
            for doc in filtered:
                doc_id = doc["doc_id"]
                col_card, col_del = st.columns([6, 1])
                with col_card:
                    st.markdown(_doc_card(doc), unsafe_allow_html=True)
                with col_del:
                    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)
                    if st.button("🗑", key=f"del_{doc_id}",
                                  help=f"Hapus {doc.get('name','')}",
                                  use_container_width=True):
                        with st.spinner("Menghapus..."):
                            engine.delete_document(doc_id)
                        st.success("Dihapus")
                        st.cache_resource.clear()
                        st.rerun()

            # Export
            st.divider()
            df_d = pd.DataFrame([{
                "Nama File":  d.get("name",""), "Kategori": d.get("category",""),
                "Deskripsi":  d.get("description",""), "Halaman": d.get("pages",0),
                "Chunks":     d.get("n_chunks",0), "Hal OCR": d.get("n_ocr",0),
                "Diupload":   d.get("uploaded_by",""), "Tanggal": str(d.get("created_at",""))[:10],
            } for d in docs])
            buf = BytesIO()
            df_d.to_excel(buf, index=False, engine="openpyxl")
            st.download_button("📥 Export Daftar (.xlsx)",
                                data=buf.getvalue(),
                                file_name="daftar_dokumen.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key="dl_d")

    # ────────────────────────────────────────────────────────
    # TAB 5 — SETUP
    # ────────────────────────────────────────────────────────
    with tab_setup:
        st.markdown(_section_header("⚙️", "Setup & Konfigurasi"), unsafe_allow_html=True)

        # DB status card
        if db_ok:
            st.markdown("""
<div style="background:rgba(35,134,54,0.1);border:1px solid rgba(35,134,54,0.3);
     border-radius:12px;padding:16px 20px;margin-bottom:16px;">
    <div style="font-size:0.95rem;font-weight:600;color:#2ea043;margin-bottom:4px;">
        🟢 Supabase PostgreSQL — Terhubung
    </div>
    <div style="font-size:0.8rem;color:#8b949e;">
        Database aktif. Semua dokumen dan index tersimpan permanen.
    </div>
</div>""", unsafe_allow_html=True)
        else:
            st.markdown("""
<div style="background:rgba(248,81,73,0.1);border:1px solid rgba(248,81,73,0.3);
     border-radius:12px;padding:16px 20px;margin-bottom:16px;">
    <div style="font-size:0.95rem;font-weight:600;color:#f85149;margin-bottom:4px;">
        🔴 Database Tidak Terhubung
    </div>
    <div style="font-size:0.8rem;color:#8b949e;">
        Tambahkan DATABASE_URL di Streamlit Cloud Secrets.
    </div>
</div>""", unsafe_allow_html=True)
            st.markdown("#### 🗄️ Setup Supabase")
            st.markdown("""
**Langkah (5 menit):**
1. Buka [supabase.com](https://supabase.com) → login → **New project**
2. Tunggu project ready
3. **Settings ⚙️ → Database → Connection string → URI** → Copy
4. Streamlit Cloud → app → **Settings → Secrets**:
```toml
DATABASE_URL = "postgresql://postgres.xxx:PASSWORD@aws-X.pooler.supabase.com:6543/postgres"

[anthropic]
api_key = "sk-ant-..."
```
5. Save → **Reboot app** ✅
""")

        st.divider()
        st.markdown("#### 📊 Status Sistem")
        sys_data = {
            "Versi Aplikasi":  APP_VERSION,
            "LLM Model":       "claude-sonnet-4-5",
            "OCR Model":       "claude-haiku-4-5 (Vision)",
            "Embedding":       EMBED_MODEL,
            "Database":        "✅ Terhubung" if db_ok else "❌ Tidak terhubung",
            "Total Dokumen":   str(len(docs)),
            "Total Chunks":    f"{engine.n_chunks:,}",
            "OCR Chunks":      f"{engine.n_ocr_chunks:,}",
        }
        st.table(pd.DataFrame(
            {"Parameter": list(sys_data.keys()),
             "Nilai":     list(sys_data.values())}
        ))

        st.divider()
        st.markdown("#### 📖 Panduan Penggunaan")
        st.markdown("""
| Fitur | Cara Pakai |
|-------|-----------|
| 💬 **Tanya AI** | Chat natural — gunakan nama teknis spesifik untuk hasil akurat |
| ⚡ **Shortcut** | Klik untuk tampilkan semua parameter Boiler / Turbin / Generator sekaligus |
| 📤 **Upload** | Pilih kategori yang tepat — mempengaruhi akurasi pencarian |
| 🔍 **Pencarian** | Semantik untuk pertanyaan umum, Keyword untuk nama parameter spesifik |
| 📑 **Dokumen** | Lihat semua dokumen, hapus yang tidak perlu, export ke Excel |

**Tips hasil terbaik:**
- Gunakan nama parameter teknis: `thrust pad temperature`, `lube oil pressure`
- Pilih filter kategori untuk membatasi ke topik tertentu
- Upload dokumen lengkap (termasuk cover dan daftar isi) untuk konteks lebih baik
""")


# ── Entry point ─────────────────────────────────────────────────
if __name__ == "__main__":
    main()
